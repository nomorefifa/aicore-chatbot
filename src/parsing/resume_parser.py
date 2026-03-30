"""
강사 이력서 파싱 파이프라인

흐름: .docx 텍스트 추출 → Gemini 구조화 (Pydantic) → 섹션별 청킹
"""

import json
import logging
from pathlib import Path

from docx import Document
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from tqdm import tqdm

from .models import Chunk, ResumeData

load_dotenv()

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

PARSE_PROMPT = """아래는 강사 이력서 원문 텍스트입니다.
내용을 분석하여 주어진 스키마에 맞게 구조화하세요.
형식이 불규칙하거나 항목이 누락된 경우 판단하여 최대한 채워 넣고, 알 수 없는 항목은 null로 두세요.

=== 이력서 텍스트 ===
{resume_text}
"""


class ResumeParser:
    """
    .docx 이력서 파일을 텍스트 추출 → Gemini 구조화 → ResumeData 반환
    """

    def __init__(self, model: str = "gemini-2.5-flash"):
        llm = ChatGoogleGenerativeAI(model=model, temperature=0)
        self.structured_llm = llm.with_structured_output(ResumeData)

    def extract_text(self, file_path: str | Path) -> str:
        """
        .docx에서 단락과 표를 추출하여 하나의 텍스트로 반환.
        표는 '헤더1 | 헤더2 | ...' 형태로 행 단위 변환.
        """
        doc = Document(str(file_path))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    def parse(self, file_path: str | Path) -> ResumeData | None:
        """
        단일 .docx 파일을 파싱하여 ResumeData 반환.
        실패 시 None 반환 후 로깅.
        """
        file_path = Path(file_path)
        try:
            raw_text = self.extract_text(file_path)
            prompt = PARSE_PROMPT.format(resume_text=raw_text)
            result: ResumeData = self.structured_llm.invoke(prompt)
            logger.info(f"파싱 완료: {file_path.name} → {result.instructor_name}")
            return result
        except Exception as e:
            logger.error(f"파싱 실패: {file_path.name} → {e}")
            return None

    def parse_all(self, raw_dir: str | Path) -> list[tuple[Path, ResumeData]]:
        """
        디렉토리 내 모든 .docx 파일을 일괄 파싱.
        반환: [(파일경로, ResumeData), ...] (실패 파일 제외)
        """
        raw_dir = Path(raw_dir)
        files = list(raw_dir.glob("*.docx"))
        logger.info(f"총 {len(files)}개 파일 파싱 시작")

        results = []
        for file_path in tqdm(files, desc="이력서 파싱"):
            resume = self.parse(file_path)
            if resume is not None:
                results.append((file_path, resume))

        logger.info(f"파싱 완료: {len(results)}/{len(files)}개 성공")
        return results


class ResumeChunker:
    """
    ResumeData를 섹션별 청크로 분할.

    청킹 전략:
    - 프로필 청크: 강사명 + 전문분야 + 연락처 (개요 검색용)
    - 학력 청크: 학력 전체를 하나의 청크로
    - 경력 청크: 경력 전체를 하나의 청크로
    - 강의이력 청크: 항목당 1개 (가장 세분화 - 핵심 검색 대상)
    - 자격증 청크: 자격증 전체를 하나의 청크로

    각 청크 content 형식: "강사: {name} | 섹션: {section} | {내용}"
    → 임베딩 시 컨텍스트를 포함하여 검색 정확도 향상
    """

    def chunk(self, resume: ResumeData, file_name: str) -> list[Chunk]:
        chunks = []
        base_meta = {
            "instructor_name": resume.instructor_name,
            "file_name": file_name,
            "doc_type": "강사이력서",
        }

        # 1. 프로필 청크 (개요)
        profile_parts = [f"강사명: {resume.instructor_name}"]
        if resume.phone:
            profile_parts.append(f"연락처: {resume.phone}")
        if resume.email:
            profile_parts.append(f"이메일: {resume.email}")
        if resume.expertise:
            profile_parts.append(f"전문분야: {', '.join(resume.expertise)}")
        if resume.summary:
            profile_parts.append(f"소개: {resume.summary}")

        chunks.append(Chunk(
            content=self._build_content(resume.instructor_name, "프로필", " | ".join(profile_parts)),
            metadata={**base_meta, "section": "프로필"},
        ))

        # 2. 학력 청크
        if resume.education:
            edu_lines = []
            for edu in resume.education:
                parts = [edu.school]
                if edu.major:
                    parts.append(edu.major)
                if edu.degree:
                    parts.append(edu.degree)
                if edu.graduation_year:
                    parts.append(edu.graduation_year)
                edu_lines.append(" / ".join(parts))

            chunks.append(Chunk(
                content=self._build_content(resume.instructor_name, "학력", "\n".join(edu_lines)),
                metadata={**base_meta, "section": "학력"},
            ))

        # 3. 경력 청크
        if resume.career:
            career_lines = []
            for c in resume.career:
                parts = [c.organization]
                if c.position:
                    parts.append(c.position)
                if c.period:
                    parts.append(c.period)
                if c.description:
                    parts.append(c.description)
                career_lines.append(" / ".join(parts))

            chunks.append(Chunk(
                content=self._build_content(resume.instructor_name, "경력", "\n".join(career_lines)),
                metadata={**base_meta, "section": "경력"},
            ))

        # 4. 강의이력 청크 (항목당 1개 - 핵심 검색 대상)
        for i, t in enumerate(resume.teaching_history):
            parts = [t.organization]
            if t.course_name:
                parts.append(t.course_name)
            if t.period:
                parts.append(t.period)
            if t.hours:
                parts.append(f"{t.hours} 시간")
            if t.description:
                parts.append(t.description)

            chunks.append(Chunk(
                content=self._build_content(resume.instructor_name, "강의이력", " / ".join(parts)),
                metadata={
                    **base_meta,
                    "section": "강의이력",
                    "teaching_organization": t.organization,
                    "course_name": t.course_name or "",
                    "teaching_index": i,
                },
            ))

        # 5. 자격증 청크
        if resume.certifications:
            cert_lines = []
            for cert in resume.certifications:
                parts = [cert.name]
                if cert.issuer:
                    parts.append(cert.issuer)
                if cert.date:
                    parts.append(cert.date)
                cert_lines.append(" / ".join(parts))

            chunks.append(Chunk(
                content=self._build_content(resume.instructor_name, "자격증", "\n".join(cert_lines)),
                metadata={**base_meta, "section": "자격증"},
            ))

        return chunks

    def _build_content(self, name: str, section: str, body: str) -> str:
        """임베딩용 컨텍스트 prefix 포함 텍스트 생성"""
        return f"강사: {name} | 섹션: {section} | {body}"


def load_chunks_from_parsed(parsed_dir: str | Path = "data/parsed") -> list[Chunk]:
    """
    이미 파싱된 JSON 파일에서 청크 생성 (Gemini 재호출 없음).
    파싱이 완료된 이후 임베딩만 다시 실행할 때 사용.
    """
    parsed_dir = Path(parsed_dir)
    chunker = ResumeChunker()
    all_chunks: list[Chunk] = []

    json_files = list(parsed_dir.glob("*.json"))
    logger.info(f"JSON 파일 {len(json_files)}개에서 청크 생성 시작")

    for json_path in json_files:
        try:
            resume = ResumeData.model_validate_json(json_path.read_text(encoding="utf-8"))
            chunks = chunker.chunk(resume, file_name=json_path.name)
            all_chunks.extend(chunks)
            logger.info(f"{resume.instructor_name}: {len(chunks)}개 청크")
        except Exception as e:
            logger.error(f"청크 생성 실패: {json_path.name} → {e}")

    logger.info(f"총 {len(all_chunks)}개 청크 생성 완료")
    return all_chunks


def run_pipeline(
    raw_dir: str | Path = "data/raw",
    parsed_dir: str | Path = "data/parsed",
    model: str = "gemini-2.5-flash",
) -> list[Chunk]:
    """
    전체 파이프라인 실행:
    data/raw/*.docx → 파싱 → 청킹 → data/parsed/{강사명}.json 저장 → 청크 목록 반환

    Returns:
        모든 이력서의 청크 목록 (임베딩 단계로 바로 전달 가능)
    """
    raw_dir = Path(raw_dir)
    parsed_dir = Path(parsed_dir)
    parsed_dir.mkdir(parents=True, exist_ok=True)

    parser = ResumeParser(model=model)
    chunker = ResumeChunker()

    parse_results = parser.parse_all(raw_dir)

    all_chunks: list[Chunk] = []
    for file_path, resume in parse_results:
        # 파싱 결과 JSON 저장
        output_path = parsed_dir / f"{resume.instructor_name}.json"
        output_path.write_text(
            resume.model_dump_json(indent=2, ensure_ascii=False),
            encoding="utf-8",
        )

        # 청킹
        chunks = chunker.chunk(resume, file_name=file_path.name)
        all_chunks.extend(chunks)
        logger.info(f"{resume.instructor_name}: {len(chunks)}개 청크 생성")

    logger.info(f"파이프라인 완료 | 총 {len(all_chunks)}개 청크")
    return all_chunks
