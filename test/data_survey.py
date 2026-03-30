# notebooks/01_data_survey.ipynb

import os
from collections import Counter

raw_dir = "./data/raw"
files = os.listdir(raw_dir)

# 파일 형식 분포
extensions = Counter(os.path.splitext(f)[1].lower() for f in files)
print(f"총 파일 수: {len(files)}")
print(f"형식 분포: {dict(extensions)}")

# DOC/DOCX 파일 하나 열어서 텍스트 추출 테스트
from docx import Document

sample = [f for f in files if f.endswith('.docx')][0]
doc = Document(os.path.join(raw_dir, sample))

# 일반 텍스트 (paragraph)
for para in doc.paragraphs[:10]:
    if para.text.strip():
        print(f"[문단] {para.text}")

# 표 내용
for i, table in enumerate(doc.tables):
    print(f"\n[표 {i+1}] {len(table.rows)}행 x {len(table.columns)}열")
    for row in table.rows[:3]:  # 앞 3행만
        print([cell.text.strip() for cell in row.cells])