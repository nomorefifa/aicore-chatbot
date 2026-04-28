"""
DOCX 텍스트 추출

python-docx를 사용해 단락과 표를 추출.
실패 시 ZIP에서 document.xml 직접 파싱하는 폴백 로직 포함.
텍스트 추출만 담당. 구조화는 DocumentParser(base_parser.py)가 처리.
"""

import logging
import re
import zipfile
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)


def _extract_text_fallback(file_path: Path) -> str:
    """python-docx 실패 시 ZIP에서 document.xml을 직접 읽어 텍스트 추출."""
    with zipfile.ZipFile(str(file_path), 'r') as z:
        xml_candidates = [n for n in z.namelist() if n.endswith('.xml') and 'word/' in n]
        doc_xml = None
        for candidate in ['word/document.xml', 'word/document2.xml']:
            if candidate in z.namelist():
                doc_xml = z.read(candidate).decode('utf-8', errors='ignore')
                break
        if doc_xml is None and xml_candidates:
            doc_xml = z.read(xml_candidates[0]).decode('utf-8', errors='ignore')

    if not doc_xml:
        return ""

    text = re.sub(r'<w:t[^>]*>', '\n', doc_xml)
    text = re.sub(r'<[^>]+>', '', text)
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    return "\n".join(lines)


def extract_text(file_path: str | Path) -> str:
    """DOCX에서 단락과 표를 추출."""
    file_path = Path(file_path)
    try:
        doc = Document(str(file_path))
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        logger.warning(f"python-docx 실패, ZIP 폴백 시도: {file_path.name} ({e})")
        return _extract_text_fallback(file_path)
