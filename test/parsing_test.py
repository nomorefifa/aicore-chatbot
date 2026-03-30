from dotenv import load_dotenv
load_dotenv()
from docx import Document
from langchain_google_genai import ChatGoogleGenerativeAI
import json
import re

# 1. 이력서 텍스트 추출 함수
def extract_text_from_docx(filepath):
    doc = Document(filepath)
    text_parts = []
    
    # 문단 텍스트
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # 표 텍스트 (행 단위로)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    
    return "\n".join(text_parts)

# 2. Gemini로 구조화
llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0)

prompt = """아래 강사 이력서 텍스트를 분석하여 다음 JSON 형식으로 구조화해주세요.
형식이 정형화되어 있지 않을 수 있으니, 내용을 기반으로 최대한 판단해주세요.
JSON만 출력하고 다른 텍스트는 포함하지 마세요.

{{
  "강사명": "",
  "전화번호": "",
  "학력": [{{"학교": "", "전공": "", "학위": "", "졸업년도": ""}}],
  "경력": [{{"기관": "", "직위": "", "기간": "", "내용": ""}}],
  "자격증": [{{"자격증명": "", "발급기관": "", "취득일": ""}}],
  "강의이력": [{{"기관": "", "과정명": "", "기간": "", "시수": "", "내용": ""}}],
  "전문분야": [],
  "기타": ""
}}

=== 이력서 텍스트 ===
{resume_text}
"""

# 3. 테스트 실행
sample_path = "data/raw/강사 프로필_박영준.docx"
raw_text = extract_text_from_docx(sample_path)
print("=== 추출된 원본 텍스트 ===")
print(raw_text[:500])

response = llm.invoke(prompt.format(resume_text=raw_text))
print("\n=== Gemini 구조화 결과 ===")
print(response.content)

# JSON 파싱 확인
clean_json = re.sub(r'^```json\s*|\s*```$', '', response.content.strip())
parsed = json.loads(clean_json)
print(f"\n강사명: {parsed['강사명']}")
print(f"강의이력 수: {len(parsed['강의이력'])}")